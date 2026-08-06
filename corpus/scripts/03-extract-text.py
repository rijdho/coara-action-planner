#!/usr/bin/env python3
"""
03-extract-text — turn the downloaded documents into plain text.

For every <work>/plans/<id>/ folder: run `pdftotext` over each PDF, read each DOCX,
concatenate them into <id>/text.txt, and finally write <work>/corpus.jsonl — one JSON
object per plan, merging _meta.json with the extracted `text` and a `chars` count.

That JSONL is what 04-theme-frequency.py counts over, and it is the only place the
plans' full text is needed. Nothing downstream sends text to a model or a network
service: the measurement is regular expressions over local files.

Incremental by default. A plan that already has a text.txt is left alone, so a
quarterly update re-extracts only what 01-harvest just downloaded — but corpus.jsonl
is always rebuilt from every folder present, because the frequencies must be counted
over the whole corpus, not the increment.

Requires: poppler's `pdftotext` on PATH. `python-docx` only if the corpus contains
DOCX files; without it those are recorded as unextractable rather than crashing.

Usage:
    python3 scripts/03-extract-text.py [--work work] [--force]
"""

import argparse
import json
import subprocess
import sys
from pathlib import Path

try:
    import docx  # python-docx
except ImportError:
    docx = None

ROOT = Path(__file__).resolve().parent.parent
MIN_USABLE_CHARS = 200


def extract_pdf(path):
    """Text from a PDF via poppler. No -layout: flow order reads better for prose."""
    try:
        out = subprocess.run(["pdftotext", "-q", str(path), "-"], capture_output=True, timeout=120)
        return out.stdout.decode("utf-8", errors="replace")
    except FileNotFoundError:
        sys.exit("pdftotext not found — install poppler (macOS: brew install poppler)")
    except Exception as exc:  # noqa: BLE001
        return f"[pdftotext failed: {exc}]"


def extract_docx(path):
    if docx is None:
        return "[python-docx not installed]"
    try:
        d = docx.Document(str(path))
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                parts.append("\t".join(c.text for c in row.cells))
        return "\n".join(parts)
    except Exception as exc:  # noqa: BLE001
        return f"[docx failed: {exc}]"


def extract_folder(folder):
    chunks = []
    for f in sorted(folder.iterdir()):
        suffix = f.suffix.lower()
        if suffix == ".pdf":
            chunks.append(extract_pdf(f))
        elif suffix == ".docx":
            chunks.append(extract_docx(f))
    return "\n\n".join(c for c in chunks if c.strip())


def main():
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--work", default="work")
    ap.add_argument("--force", action="store_true", help="re-extract plans that already have text.txt")
    args = ap.parse_args()

    work = ROOT / args.work
    plans = work / "plans"
    if not plans.is_dir():
        sys.exit(f"{plans} not found — run: node scripts/01-harvest.mjs")

    folders = sorted((p for p in plans.iterdir() if p.is_dir() and p.name.isdigit()),
                     key=lambda p: int(p.name))
    if not folders:
        sys.exit(f"no plan folders in {plans}")

    records, thin, reused = [], [], 0
    for i, folder in enumerate(folders, 1):
        text_path = folder / "text.txt"
        if text_path.exists() and not args.force:
            text = text_path.read_text(encoding="utf-8")
            reused += 1
        else:
            text = extract_folder(folder)
            text_path.write_text(text, encoding="utf-8")

        meta_path = folder / "_meta.json"
        meta = json.loads(meta_path.read_text(encoding="utf-8")) if meta_path.exists() else {}
        records.append({
            "id": meta.get("id", int(folder.name)),
            "doi": meta.get("doi"),
            "title": meta.get("title"),
            "creators": meta.get("creators", []),
            "date": meta.get("date"),
            "language": meta.get("language"),
            "chars": len(text),
            "text": text,
        })
        if len(text.strip()) < MIN_USABLE_CHARS:
            thin.append(folder.name)
        if i % 25 == 0 or i == len(folders):
            print(f"  [{i}/{len(folders)}] {folder.name} — {len(text):,} chars")

    corpus = work / "corpus.jsonl"
    with corpus.open("w", encoding="utf-8") as fh:
        for rec in records:
            fh.write(json.dumps(rec, ensure_ascii=False) + "\n")

    total = sum(r["chars"] for r in records)
    print(f"\n{len(records)} plans, {total:,} chars ({reused} reused, {len(records) - reused} extracted).")
    print(f"Corpus: {corpus}")
    if thin:
        # These still count in the denominator: a plan whose text could not be read is
        # a plan that mentions nothing, which can only push prevalence down. Dropping
        # them would quietly inflate every percentage.
        print(f"\n{len(thin)} plan(s) under {MIN_USABLE_CHARS} chars — probably scans needing OCR: "
              f"{', '.join(thin)}")
    print("\nNext: python3 scripts/04-theme-frequency.py " + str(corpus))


if __name__ == "__main__":
    main()
